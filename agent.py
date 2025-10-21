# # import io
# # import docx
# # from typing import List
# # from openai import OpenAI, AsyncOpenAI
# # from langchain_openai import OpenAIEmbeddings, ChatOpenAI
# # from langchain.text_splitter import RecursiveCharacterTextSplitter
# # from langchain_community.vectorstores import FAISS
# # from langchain.chains import create_retrieval_chain
# # from langchain.chains.combine_documents import create_stuff_documents_chain
# # from langchain_core.prompts import ChatPromptTemplate
# # from langchain_core.documents import Document
# # from fastapi import UploadFile

# # class ReportAgent:
# #     """Handles report generation and Q&A using a RAG pipeline."""

# #     def __init__(self, openai_api_key: str):
# #         """Initializes the agent with OpenAI clients and necessary components."""
# #         if not openai_api_key:
# #             raise ValueError("OpenAI API key is missing.")
# #         self.client = OpenAI(api_key=openai_api_key)
# #         self.async_client = AsyncOpenAI(api_key=openai_api_key)
        
# #         self.retriever = None
# #         self.qa_chain = None
# #         self.report_context = ""

# #     # ... (the _parse_docx, _create_dynamic_prompt, and stream_report_generation methods remain the same) ...
# #     def _parse_docx(self, file_bytes: bytes) -> str:
# #         """Extracts text content from .docx file bytes."""
# #         try:
# #             doc = docx.Document(io.BytesIO(file_bytes))
# #             return "\n".join([para.text for para in doc.paragraphs if para.text])
# #         except Exception as e:
# #             print(f"Error parsing DOCX file: {e}")
# #             return ""

# #     def _create_dynamic_prompt(self, blueprint_text: str) -> str:
# #         """
# #         Analyzes the blueprint text and generates a dynamic, optimized summary prompt.
# #         """
# #         meta_prompt = (
# #             "You are an expert prompt engineering assistant. I will provide you with the text from a document that "
# #             "contains instructions, a template, or a blueprint for a task. Your job is to analyze this text and create "
# #             "a concise, powerful, and direct summary prompt for another AI to perform the main task."
# #             "\n\nThe summary prompt you create must:"
# #             "\n1. Start by assigning a clear **role** to the AI (e.g., 'Act as a...')."
# #             "\n2. State the primary **goal** or objective of the task."
# #             "\n3. Identify the main **subject** or target (e.g., a specific country, topic, or product)."
# #             "\n4. Summarize the 3-4 most **critical instructions**, constraints, or formatting rules mentioned in the document."
# #             "\n5. End with a clear **call to action** to begin the task immediately."
# #             "\n\nHere is the document text:"
# #             f"\n\n---\n{blueprint_text}\n---"
# #         )
        
# #         try:
# #             response = self.client.chat.completions.create(
# #                 model="gpt-4",
# #                 messages=[
# #                     {"role": "system", "content": "You are a helpful prompt engineering assistant."},
# #                     {"role": "user", "content": meta_prompt}
# #                 ],
# #                 temperature=0.2,
# #             )
# #             dynamic_prompt = response.choices[0].message.content
# #             print(f"✅ Dynamically generated prompt:\n{dynamic_prompt}")
# #             return dynamic_prompt
# #         except Exception as e:
# #             print(f"⚠️ Error creating dynamic prompt: {e}. Falling back to a generic prompt.")
# #             return "Based on the provided blueprint, generate a comprehensive report."


# #     async def stream_report_generation(self, files: List[UploadFile], user_prompt: str):
# #         """
# #         Generates a report from one or more .docx blueprints using a streaming response.
# #         """
# #         combined_blueprint_text = ""
# #         for i, file in enumerate(files):
# #             file_bytes = await file.read()
# #             text = self._parse_docx(file_bytes)
# #             combined_blueprint_text += f"--- CONTENT FROM FILE {i+1}: {file.filename} ---\n"
# #             combined_blueprint_text += text + "\n\n"

# #         if user_prompt:
# #             final_user_prompt = user_prompt
# #         else:
# #             final_user_prompt = self._create_dynamic_prompt(combined_blueprint_text)

# #         system_prompt = (
# #             "You are an expert report writer. You will be given a blueprint and a summary of your task. "
# #             "Follow the instructions precisely to generate a comprehensive and well-structured report in Markdown format. "
# #             "Fill in any placeholders like '[...]' with relevant, detailed information."
# #         )
        
# #         full_prompt_for_generation = (
# #             f"**TASK SUMMARY:**\n{final_user_prompt}\n\n"
# #             f"**FULL BLUEPRINT DOCUMENT(S):**\n{combined_blueprint_text}"
# #         )

# #         stream = await self.async_client.chat.completions.create(
# #             model="gpt-4-turbo",
# #             messages=[
# #                 {"role": "system", "content": system_prompt},
# #                 {"role": "user", "content": full_prompt_for_generation},
# #             ],
# #             stream=True,
# #         )

# #         full_report = ""
# #         async for chunk in stream:
# #             content = chunk.choices[0].delta.content
# #             if content:
# #                 full_report += content
# #                 yield content
        
# #         self.report_context = full_report
# #         self.setup_rag_pipeline(self.report_context)

# #     def setup_rag_pipeline(self, report_markdown: str):
# #         if not report_markdown:
# #             print("⚠️ Warning: Report content is empty. RAG pipeline not set up.")
# #             return

# #         text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
# #         docs = [Document(page_content=report_markdown)]
# #         chunks = text_splitter.split_documents(docs)

# #         embeddings = OpenAIEmbeddings(openai_api_key=self.client.api_key)
# #         vectorstore = FAISS.from_documents(documents=chunks, embedding=embeddings)
# #         self.retriever = vectorstore.as_retriever(search_kwargs={"k": 3})

# #         system_prompt = (
# #             "You are an assistant for question-answering tasks. Use the following retrieved "
# #             "context to answer the question. If you don't know the answer, just say that you don't know. "
# #             "Use three sentences maximum and keep the answer concise.\n\n"
# #             "{context}"
# #         )
# #         prompt = ChatPromptTemplate.from_messages(
# #             [("system", system_prompt), ("human", "{input}")]
# #         )
        
# #         llm = ChatOpenAI(model_name="gpt-4-turbo", temperature=0, openai_api_key=self.client.api_key, streaming=True)

# #         question_answer_chain = create_stuff_documents_chain(llm, prompt)
# #         self.qa_chain = create_retrieval_chain(self.retriever, question_answer_chain)
# #         print("✅ RAG pipeline is ready.")

# #     # --- THIS METHOD IS NOW A STREAMING GENERATOR ---
# #     async def stream_answer(self, question: str):
# #         """
# #         Streams an answer based on the generated report using the RAG pipeline.
# #         """
# #         if not self.qa_chain:
# #             yield "The RAG pipeline has not been set up. Please generate a report first."
# #             return

# #         # Use .astream() for asynchronous streaming with LangChain
# #         async for chunk in self.qa_chain.astream({"input": question}):
# #             if answer_chunk := chunk.get("answer"):
# #                 yield answer_chunk

# import io
# import docx
# from typing import List
# from openai import OpenAI, AsyncOpenAI
# from langchain_openai import OpenAIEmbeddings, ChatOpenAI
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import FAISS
# from langchain.chains.combine_documents import create_stuff_documents_chain
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_core.documents import Document
# from langchain_core.runnables import RunnablePassthrough
# from fastapi import UploadFile

# class ReportAgent:
#     """Handles report generation and Q&A using a RAG pipeline."""

#     def __init__(self, openai_api_key: str):
#         if not openai_api_key:
#             raise ValueError("OpenAI API key is missing.")
#         self.client = OpenAI(api_key=openai_api_key)
#         self.async_client = AsyncOpenAI(api_key=openai_api_key)
        
#         self.retriever = None
#         # This will now store the specific Q&A part of the chain
#         self.question_answer_chain = None
#         self.report_context = ""

#     def _parse_docx(self, file_bytes: bytes) -> str:
#         try:
#             doc = docx.Document(io.BytesIO(file_bytes))
#             return "\n".join([para.text for para in doc.paragraphs if para.text])
#         except Exception as e:
#             print(f"Error parsing DOCX file: {e}")
#             return ""

#     def _create_dynamic_prompt(self, blueprint_text: str) -> str:
#         meta_prompt = (
#             "You are an expert prompt engineering assistant. I will provide you with the text from a document that "
#             "contains instructions, a template, or a blueprint for a task. Your job is to analyze this text and create "
#             "a concise, powerful, and direct summary prompt for another AI to perform the main task."
#             "\n\nThe summary prompt you create must:"
#             "\n1. Start by assigning a clear **role** to the AI (e.g., 'Act as a...')."
#             "\n2. State the primary **goal** or objective of the task."
#             "\n3. Identify the main **subject** or target (e.g., a specific country, topic, or product)."
#             "\n4. Summarize the 3-4 most **critical instructions**, constraints, or formatting rules mentioned in the document."
#             "\n5. End with a clear **call to action** to begin the task immediately."
#             "\n\nHere is the document text:"
#             f"\n\n---\n{blueprint_text}\n---"
#         )
#         try:
#             response = self.client.chat.completions.create(
#                 model="gpt-4o-mini",
#                 messages=[{"role": "system", "content": "You are a helpful prompt engineering assistant."}, {"role": "user", "content": meta_prompt}],
#                 temperature=0.2,
#             )
#             return response.choices[0].message.content
#         except Exception as e:
#             print(f"⚠️ Error creating dynamic prompt: {e}. Falling back to a generic prompt.")
#             return "Based on the provided blueprint, generate a comprehensive report."

#     async def stream_report_generation(self, files: List[UploadFile], user_prompt: str):
#         combined_blueprint_text = ""
#         for i, file in enumerate(files):
#             file_bytes = await file.read()
#             text = self._parse_docx(file_bytes)
#             combined_blueprint_text += f"--- CONTENT FROM FILE {i+1}: {file.filename} ---\n{text}\n\n"

#         final_user_prompt = user_prompt or self._create_dynamic_prompt(combined_blueprint_text)
#         system_prompt = (
#             "You are an expert report writer. You will be given a blueprint and a summary of your task. "
#             "Follow the instructions precisely to generate a comprehensive and well-structured report in Markdown format. "
#             "Fill in any placeholders like '[...]' with relevant, detailed information."
#         )
#         full_prompt_for_generation = (
#             f"**TASK SUMMARY:**\n{final_user_prompt}\n\n**FULL BLUEPRINT DOCUMENT(S):**\n{combined_blueprint_text}"
#         )

#         stream = await self.async_client.chat.completions.create(
#             model="gpt-4-turbo",
#             messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": full_prompt_for_generation}],
#             stream=True,
#         )

#         full_report = ""
#         async for chunk in stream:
#             content = chunk.choices[0].delta.content
#             if content:
#                 full_report += content
#                 yield content
        
#         self.report_context = full_report
#         self.setup_rag_pipeline(self.report_context)

#     def setup_rag_pipeline(self, report_markdown: str):
#         if not report_markdown:
#             print("⚠️ Warning: Report content is empty. RAG pipeline not set up.")
#             return

#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
#         docs = [Document(page_content=report_markdown)]
#         chunks = text_splitter.split_documents(docs)

#         embeddings = OpenAIEmbeddings(openai_api_key=self.client.api_key)
#         vectorstore = FAISS.from_documents(documents=chunks, embedding=embeddings)
#         self.retriever = vectorstore.as_retriever(search_kwargs={"k": 3})

#         system_prompt = (
#             "You are an assistant for question-answering tasks. Use the following retrieved "
#             "context to answer the question. If you don't know the answer, just say that you don't know. "
#             "Keep the answer concise.\n\n{context}"
#         )
#         prompt = ChatPromptTemplate.from_messages([("system", system_prompt), ("human", "{input}")])
#         llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0, openai_api_key=self.client.api_key, streaming=True)
        
#         # We now store the specific document-handling chain
#         self.question_answer_chain = create_stuff_documents_chain(llm, prompt)
#         print("✅ RAG pipeline is ready.")

#     # --- THIS IS THE CORRECTED STREAMING METHOD ---
#     async def stream_answer(self, question: str):
#         """
#         Streams an answer by creating a new, stable RAG chain on the fly.
#         """
#         if not self.question_answer_chain or not self.retriever:
#             yield "The RAG pipeline has not been set up. Please generate a report first."
#             return

#         # This chain correctly pipes the question to the retriever and then passes
#         # both the question and the retrieved context to the final QA chain.
#         rag_chain = (
#             RunnablePassthrough.assign(context=(lambda x: x["input"]) | self.retriever)
#             | self.question_answer_chain
#         )

#         # Stream the response from this stable chain
#         async for chunk in rag_chain.astream({"input": question}):
#             yield chunk



# import io
# import docx
# import re
# from typing import List
# from openai import OpenAI, AsyncOpenAI
# from langchain_openai import OpenAIEmbeddings, ChatOpenAI
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import FAISS
# from langchain.chains.combine_documents import create_stuff_documents_chain
# from langchain_core.prompts import ChatPromptTemplate
# from langchain_core.documents import Document
# from langchain_core.runnables import RunnablePassthrough
# from fastapi import UploadFile

# class ReportAgent:
#     """Handles report generation and Q&A using a RAG pipeline."""

#     def __init__(self, openai_api_key: str):
#         if not openai_api_key:
#             raise ValueError("OpenAI API key is missing.")
#         self.client = OpenAI(api_key=openai_api_key)
#         self.async_client = AsyncOpenAI(api_key=openai_api_key)
        
#         self.retriever = None
#         self.question_answer_chain = None
#         self.report_context = ""

#     def _parse_docx(self, file_bytes: bytes) -> str:
#         try:
#             doc = docx.Document(io.BytesIO(file_bytes))
#             return "\n".join([para.text for para in doc.paragraphs if para.text])
#         except Exception as e:
#             print(f"Error parsing DOCX file: {e}")
#             return ""

#     def _create_dynamic_prompt(self, blueprint_text: str) -> str:
#         # This function is now less critical, but we keep it as a fallback
#         meta_prompt = (
#             "You are an expert prompt engineering assistant. Analyze the following document text and create a concise "
#             "summary prompt for another AI to perform the main task. The prompt must assign a role, state the primary goal, "
#             "identify the main subject, and list 3-4 critical instructions."
#             f"\n\n---\n{blueprint_text}\n---"
#         )
#         try:
#             response = self.client.chat.completions.create(
#                 model="gpt-4",
#                 messages=[{"role": "system", "content": "You are a helpful prompt engineering assistant."}, {"role": "user", "content": meta_prompt}],
#                 temperature=0.2,
#             )
#             return response.choices[0].message.content
#         except Exception as e:
#             print(f"⚠️ Error creating dynamic prompt: {e}. Falling back to a generic prompt.")
#             return "Based on the provided blueprint, generate a comprehensive report."

#     async def stream_report_generation(self, combined_blueprint_text: str, user_prompt: str):
#         """
#         Generates a report from the combined text of the blueprints.
#         """
        
#         # --- THIS IS THE NEW, HIGH-QUALITY SYSTEM PROMPT ---
#         system_prompt = (
#     "You are a market research assistant tasked with generating a professional, comprehensive market report for biomass and energy in the specified country/region. "
#     "Retrieve relevant data from the embedded documents and use the provided blueprint to generate the report, adhering strictly to the following structure:\n"
    
#     "1. **Headings & Subheadings**: Use **bold, clear headings** for each section and subheading (e.g., 'A. General Country Information', 'B. Biomass & Waste Management'). "
#     "Ensure consistent and professional formatting throughout the report.\n"
    
#     "2. **Data-Rich Content**: Include **quantitative data**, such as **biomass consumption, energy demand** forecasts, and **market statistics** where possible. "
#     "For each 'TOP-5' list (e.g., biomass manufacturers, transport companies), use **GitHub-flavored Markdown tables** with columns for **Name**, **Headquarters**, **Volume**, and **Contact Information**.\n"
    
#     "3. **Actionable Insights**: Provide **clear, actionable recommendations** at the end of each section. These should help guide business strategy, e.g., 'Invest in biomass processing infrastructure in Region X.'\n"
    
#     "4. **Conciseness & Clarity**: Use **bullet points**, **numbered lists**, and **short paragraphs** to ensure the report is easy to read and understand. Avoid lengthy explanations and focus on presenting key insights.\n"
    
#     "5. **Visuals**: If relevant, **insert graphs** or **charts** to represent data (e.g., biomass supply, energy trends, market growth)."
    
#     "The final report should be **professional, well-structured**, and ready to be used for business decision-making. "
#     "Generate the ful report in English."
# )


        
#         # Use the user's text-box prompt if provided, otherwise, generate one.
#         final_user_prompt = user_prompt or self._create_dynamic_prompt(combined_blueprint_text)
        
#         # Combine the user's prompt (from text box) with the full blueprint text
#         full_prompt_for_generation = (
#             f"**User's Main Request (Use this for focus, e.g., country):**\n{final_user_prompt}\n\n"
#             f"**Full Blueprint Document(S) Content (Use this for structure and content):**\n{combined_blueprint_text}"
#         )

#         stream = await self.async_client.chat.completions.create(
#             model="gpt-4-turbo",
#             messages=[
#                 {"role": "system", "content": system_prompt},
#                 {"role": "user", "content": full_prompt_for_generation}
#             ],
#             stream=True,
#         )

#         full_report = ""
#         async for chunk in stream:
#             content = chunk.choices[0].delta.content
#             if content:
#                 full_report += content
#                 yield content
        
#         self.report_context = full_report
#         self.setup_rag_pipeline(self.report_context)

#     def setup_rag_pipeline(self, report_markdown: str):
#         """Sets up the RAG pipeline by creating a FAISS vector store."""
#         if not report_markdown:
#             print("⚠️ Warning: Report content is empty. RAG pipeline not set up.")
#             return

#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
#         docs = [Document(page_content=report_markdown)]
#         chunks = text_splitter.split_documents(docs)

#         embeddings = OpenAIEmbeddings(openai_api_key=self.client.api_key)
#         vectorstore = FAISS.from_documents(documents=chunks, embedding=embeddings)
#         self.retriever = vectorstore.as_retriever(search_kwargs={"k": 3})

#         system_prompt = (
#             "You are an assistant for question-answering tasks. Use the following retrieved "
#             "context to answer the question. If you don't know the answer, just say that you don't know. "
#             "Keep the answer concise.\n\n{context}"
#         )
#         prompt = ChatPromptTemplate.from_messages([("system", system_prompt), ("human", "{input}")])
#         llm = ChatOpenAI(model_name="gpt-4-turbo", temperature=0, openai_api_key=self.client.api_key, streaming=True)
        
#         self.question_answer_chain = create_stuff_documents_chain(llm, prompt)
#         print("✅ RAG pipeline is ready.")

#     async def stream_answer(self, question: str):
#         """Streams an answer based on the generated report using the RAG pipeline."""
#         if not self.question_answer_chain or not self.retriever:
#             yield "The RAG pipeline has not been set up. Please generate a report first."
#             return

#         rag_chain = (
#             RunnablePassthrough.assign(context=(lambda x: x["input"]) | self.retriever)
#             | self.question_answer_chain
#         )

#         async for chunk in rag_chain.astream({"input": question}):
#             yield chunk

import io
import docx
import re
from typing import List
from openai import OpenAI, AsyncOpenAI
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.documents import Document
from langchain_core.runnables import RunnablePassthrough
from fastapi import UploadFile

class ReportAgent:
    """Handles report generation and Q&A using a RAG pipeline."""

    def __init__(self, openai_api_key: str):
        if not openai_api_key:
            raise ValueError("OpenAI API key is missing.")
        self.client = OpenAI(api_key=openai_api_key)
        self.async_client = AsyncOpenAI(api_key=openai_api_key)
        self.retriever = None
        self.question_answer_chain = None
        self.report_context = ""

    def _parse_docx(self, file_bytes: bytes) -> str:
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n".join([para.text for para in doc.paragraphs if para.text])
        except Exception as e:
            print(f"Error parsing DOCX file: {e}")
            return ""

    def _create_dynamic_prompt(self, blueprint_text: str) -> str:
        # Fallback function. The main instructions are in the system prompt.
        return "Generate a comprehensive market report based on the provided blueprints."

    async def stream_report_generation(self, combined_blueprint_text: str, user_prompt: str):
        
        # --- THIS IS THE UPDATED SYSTEM PROMPT WITH MORE DETAIL ---
        # --- THIS IS THE UPDATED SYSTEM PROMPT WITH PAGE LIMIT ---
        # --- THIS IS THE UPDATED SYSTEM PROMPT FOR A DETAILED 7-8 PAGE REPORT ON ITALY ---
        # --- THIS IS THE UPDATED SYSTEM PROMPT FOR A DETAILED 7-8 PAGE REPORT ON ITALY (ENGLISH ONLY) ---
        # --- THIS IS THE UPDATED SYSTEM PROMPT FOR A DETAILED 7-8 PAGE REPORT ON GERMANY (ENGLISH ONLY, LATEST DATA) ---
        # --- THIS IS THE UPDATED SYSTEM PROMPT FOR A DETAILED 7-8 PAGE REPORT ON ITALY (GERMAN ONLY, LATEST DATA) ---
        # --- THIS IS THE UPDATED SYSTEM PROMPT FOR A DETAILED 7-8 PAGE REPORT ON GERMANY (ENGLISH ONLY, LATEST DATA) ---
        # --- THIS IS THE UPDATED SYSTEM PROMPT FOR A DETAILED 7-8 PAGE REPORT ON GERMANY (GERMAN ONLY, LATEST DATA) ---
        system_prompt = (
            "You are a top-tier market research analyst from a leading consultancy firm. Your client, a German biomass trading company, has provided detailed blueprints and an example of a high-quality report ('Markt-Report Italien'). Your task is to generate a new, comprehensive market report for **GERMANY** that meets this high standard of detail, length (approximately **7-8 pages**), professionalism, and uses the **most current data available (ideally 2024-2025)**. Failure to follow these rules is a critical error."
            "Generate Fully detailed report"
            "\n### NON-NEGOTIABLE COMMANDS ###\n"
            "1.  **LATEST DATA:** You MUST prioritize using the most recent data available, ideally from the last **12-24 months (2024-2025 preferred)**. If only older data is accessible, clearly state the year. Mark estimated figures with '(geschätzt)'."
            "2.  **TARGET LENGTH & DETAIL:** Your output MUST be detailed and comprehensive, aiming for a length of approximately **7-8 pages** when rendered as a PDF. Match the detail level found in the 'Markt-Report Italien Biomasse_GPT-5_v1.1.pdf' example (multi-paragraph sections, specific data, nuanced analysis), adapting content for Germany. Short, superficial answers are a failure."
            "3.  **COMPLETE ALL SECTIONS (A-H):** You MUST generate content for ALL sections (A through H) and all sub-points listed in the user's blueprints, applying them to the German context. This includes 'Generelle Landesinformationen', 'Energieversorgung', 'Biomasse & Abfallwirtschaft', all TOP-5s for 'Transport & Logistik', 'Regulatorik', 'Geschäftskultur', 'Strategien', and 'Fazit & Handlungsempfehlungen (H)'. DO NOT SKIP ANY."
            "4.  **STRICT MARKDOWN FORMATTING:**\n"
            "    * **Headings:** MUST use Markdown headers (`#`, `##`, `###`). DO NOT use bold text (e.g., `**A. General...**`) for headings. `#` for the main title, `##` for main sections (A, B, C...), and `###` for sub-sections.\n"
            "    * **Bolding:** Use `**bold text**` for emphasis *within* paragraphs or bullet points, but NEVER for section headings.\n"
            "5.  **DATA-RICH TABLES (CRITICAL FAILURE POINT):**\n"
            "    * All 'TOP-5' lists MUST be formatted as **strict GitHub-flavored Markdown tables.**\n"
            "    * **FORBIDDEN:** The phrase 'The following table:' (or its German equivalent 'Die folgende Tabelle:') is FORBIDDEN. Do NOT use it.\n"
            "    * **REQUIRED FORMAT:**\n"
            "        | Name | Headquarters | Volume (Annual) | Contact (Plausible) |\n"
            "        | :--- | :--- | :--- | :--- |\n"
            "        | Beispiel Firma GmbH | Berlin | 150K Tonnen | info@beispiel.de |\n"
            "    * You MUST generate plausible, realistic data relevant to Germany for ALL requested columns."
            "6.  **WORKFLOW:** You MUST generate the **full, complete report in one single response.** IGNORE any instructions about 'waiting for feedback'."
            "7.  **LANGUAGE:** You MUST generate the entire report **only in German**."
            "8.  **NO CITATIONS:** Do NOT include source citations (e.g., '') in the final output."

            "\n### REQUIRED CONTENT DETAIL (Per Section for GERMANY - Be Detailed, Use Latest Data, Write in German) ###\n"
            "**## A. Generelle Landesinformationen**:\n"
            "   - **### Geografie, Klima, Forstwirtschaft, Demografie**: Include recent economic context (BIP, Hauptindustrien wie Automobil, Chemie, Fertigung), regional differences (e.g., Nord vs. Süd, Ost vs. West). Detail forestry specifics (dominante Arten wie Fichte, Kiefer, Buche, Eiche; aktuelle Ernte vs. Zuwachsraten, Import-/Exportbilanz für Deutschland).\n"
            
            "**## B. Energieversorgung & Trends**:\n"
            "   - **### Energiemix**: Provide recent % breakdown for ALL sources in Germany (Gas, Öl, Erneuerbare - spezifizieren Sie Wind, Solar, Biomasse, Wasser, Kohle/Braunkohle, Status Kernenergieausstieg, Importe). Discuss current trends (Fortschritt Energiewende, Kohleausstiegspläne) and cite specific national energy targets (EEG, Klimaschutzgesetz). Mention current infrastructure status (Gaspipelines wie Nord Stream-Implikationen, Netzausbaubedarf).\n"
            
            "**## C. Biomasse & Abfallwirtschaft**:\n"
            "   - **### Marktspezifika**: Include recent data for Germany's wood pellet market (Produktion, Verbrauch, Preistrends, Import/Export) and for Holzchips (Verwendung in KWK-Anlagen, Hauptquellen). Identify major biomass associations (e.g., DEPV, FNR).\n"
            "   - **### Top-5 Biomasse Lieferanten (Tabelle)**: Include current Name, Sitz, Volumen, Kontakt, Zertifizierungen relevant to the German market.\n"
            "   - **### Top-5 Biomasse Verbraucher (Tabelle)**: Include current Name, Standort (e.g., große KWK-Anlagen), Volumen, Kontakt.\n"
            "   - **### Abfallwirtschaft**: Include recent recycling rates (kommunal, Verpackungen - Grüner Punkt), Deponieverbote (Kreislaufwirtschaftsgesetz), specifics on regulations like Gewerbeabfallverordnung in Germany.\n"
            
            "**## D. Transport & Logistik**:\n"
            "   - **### Logistikinfrastruktur**: Detail current status of ports (Hamburg, Bremen/Bremerhaven, Rostock), rail network (DB Cargo), key waterways (Rhine, Elbe), inland hubs relevant to biomass in Germany.\n"
            "   - **### Top-5 Logistikdienstleister (Separate Tabellen für LKW, Bahn, Schiff)**: Include current Name, Sitz, Volumen/Umsatz, Kontakt specific to Germany (e.g., major freight forwarders, rail operators).\n"
            
            "**## E. Regulatorik & Politische Landschaft**:\n"
            "   - **### Politik**: Detail current German laws/programs (EEG subsidies for biomass, Renewable Heat Act - GEG) and their incentives/tariffs.\n"
            
            "**## F. Geschäfts- & Handelskultur**:\n"
            "   - **### Kulturelle Tipps**: Actionable advice on current German business etiquette (Pünktlichkeit is crucial, formal address 'Sie', direct communication style).\n"
            "   - **### Markteintritt**: Concrete strategy recommendations (e.g., office location in Germany, partnerships with Stadtwerke or agricultural cooperatives).\n"
            
            "**## G. Expansionsstrategien & Wachstumshebel**:\n"
            "   - **### Konkrete Strategien**: Make expansion blueprints specific to Germany with suggested actions/locations (e.g., focus on regions with high forestry or CHP plant density).\n"
            
            "**## H. Fazit & Handlungsempfehlungen**:\n" 
            "   - **### Priorisierte Maßnahmen**: Provide specific, prioritized, actionable recommendations for Germany based on current market conditions (target regions, compliance needs like EEG/GEG, partnership opportunities).\n"

            "Generate the full, detailed, professional, German-only report for GERMANY now, using the latest available data (2024-2025 preferred) and aiming for approximately 7-8 pages."
        )


        final_user_prompt = user_prompt or self._create_dynamic_prompt(combined_blueprint_text)
        
        full_prompt_for_generation = (
            f"**User's Main Request (Use this for focus, e.g., country):**\n{final_user_prompt}\n\n"
            f"**Full Blueprint Document(S) Content (Use this for structure and context):**\n{combined_blueprint_text}"
        )

        stream = await self.async_client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": full_prompt_for_generation}
            ],
            stream=True,
        )

        full_report = ""
        async for chunk in stream:
            content = chunk.choices[0].delta.content
            if content:
                full_report += content
                yield content
        
        self.report_context = full_report
        self.setup_rag_pipeline(self.report_context)

    def setup_rag_pipeline(self, report_markdown: str):
        if not report_markdown:
            print("⚠️ Warning: Report content is empty. RAG pipeline not set up.")
            return
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        docs = [Document(page_content=report_markdown)]
        chunks = text_splitter.split_documents(docs)
        embeddings = OpenAIEmbeddings(openai_api_key=self.client.api_key)
        vectorstore = FAISS.from_documents(documents=chunks, embedding=embeddings)
        self.retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
        system_prompt = (
            "You are an assistant for question-answering tasks. Use the following retrieved "
            "context to answer the question. If you don't know the answer, just say that you don't know. "
            "Keep the answer concise.\n\n{context}"
        )
        prompt = ChatPromptTemplate.from_messages([("system", system_prompt), ("human", "{input}")])
        llm = ChatOpenAI(model_name="gpt-4-turbo", temperature=0, openai_api_key=self.client.api_key, streaming=True)
        self.question_answer_chain = create_stuff_documents_chain(llm, prompt)
        print("✅ RAG pipeline is ready.")

    async def stream_answer(self, question: str):
        if not self.question_answer_chain or not self.retriever:
            yield "The RAG pipeline has not been set up. Please generate a report first."
            return
        rag_chain = (
            RunnablePassthrough.assign(context=(lambda x: x["input"]) | self.retriever)
            | self.question_answer_chain
        )
        async for chunk in rag_chain.astream({"input": question}):
            yield chunk